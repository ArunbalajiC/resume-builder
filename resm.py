import streamlit as st
import docx
import PyPDF2
from duckduckgo_search import DDGS
import re
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO


def docx_to_html(docx_file):
    doc = docx.Document(docx_file)
    html = '<html><body>'
    for para in doc.paragraphs:
        html += f'<p>{para.text}</p>'
    html += '</body></html>'
    return html


# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PyPDF2.PdfReader(pdf_file)
    text = ''.join([page.extract_text() for page in pdf.pages])
    return text

# Function to set document margins
def set_margins(section, left=1, right=1, top=1, bottom=1):
    section.left_margin = docx.shared.Inches(left)
    section.right_margin = docx.shared.Inches(right)
    section.top_margin = docx.shared.Inches(top)
    section.bottom_margin = docx.shared.Inches(bottom)

# Function to generate resume content as Markdown (for preview)
def generate_resume_markdown(
    name=None,
    contact_info=None,
    professional_summary=None,
    education=None,
    experience=None,
    projects=None,
    skills=None,
    languages=None,
    links=None,
    awards=None,
    certifications=None,
    publications=None,
    volunteering=None,
    competitions=None,
    conferences_workshops=None,
    tests=None,
    patents=None,
    scholarships=None,
    extracurricular_activities=None,
):
    content = ""
    
    if name:
        content += f"# {name}\n\n"

    if contact_info:
        content += f"**Contact Information:** {contact_info}\n\n"

    if professional_summary:
        content += f"## Professional Summary\n\n{professional_summary}\n\n"

    if education:
        content += "## Education\n\n"
        for edu in education:
            content += f"**{edu['institution']}** - {edu['degree']} ({edu['cgpa']}) - {edu['dates']}\n\n"

    if experience:
        content += "## Experience\n\n"
        for exp in experience:
            content += f"**{exp['title']}** at **{exp['company']}** ({exp['dates']})\n- {exp['description']}\n\n"

    if projects:
        content += "## Projects\n\n"
        for proj in projects:
            content += f"**{proj['title']}**\n- {proj['description']}\n\n"

    if skills:
        content += "## Skills\n\n"
        for skill in skills:
            content += f"**{skill['category']}:** {', '.join(skill['skills'])}\n\n"

    if languages:
        content += "## Languages\n\n"
        content += ', '.join(languages) + "\n\n"

    if links:
        content += "## Links\n\n"
        for link in links:
            content += f"- {link}\n\n"

    if awards:
        content += "## Awards\n\n"
        for award in awards:
            content += f"- {award}\n\n"

    if certifications:
        content += "## Certifications\n\n"
        for certification in certifications:
            content += f"- {certification}\n\n"

    if publications:
        content += "## Publications\n\n"
        for publication in publications:
            content += f"- {publication}\n\n"

    if volunteering:
        content += "## Volunteering\n\n"
        for vol in volunteering:
            content += f"- {vol}\n\n"

    if competitions:
        content += "## Competitions\n\n"
        for competition in competitions:
            content += f"- {competition}\n\n"

    if conferences_workshops:
        content += "## Conferences and Workshops\n\n"
        for conference in conferences_workshops:
            content += f"- {conference}\n\n"

    if tests:
        content += "## Tests\n\n"
        for test in tests:
            content += f"- {test}\n\n"

    if patents:
        content += "## Patents\n\n"
        for patent in patents:
            content += f"- {patent}\n\n"

    if scholarships:
        content += "## Scholarships\n\n"
        for scholarship in scholarships:
            content += f"- {scholarship}\n\n"

    if extracurricular_activities:
        content += "## Extracurricular Activities\n\n"
        for activity in extracurricular_activities:
            content += f"- {activity}\n\n"

    return content

def extract_sections_from_resume(resume_text):
    part1 = """
        Extract the following details from the resume text below and start them with the exact same title as given below, and if the detail does not exist use the title and fill it with 'NA' and maintain the exact same order:

        - Name
        - Contact Information
        - Professional Summary
        - Education
        - Experience
    """
    part2 = """
        Extract the following details from the resume text below and start them with the exact same title as given below, and if the detail does not exist use the title and fill it with 'NA' and maintain the exact same order:

        - Projects
        - Skills
        - Languages
        - Links
        - Awards
        - Certifications
        - Publications
    """
    part3 = """
        Extract the following details from the resume text below and start them with the exact same title as given below, and if the detail does not exist use the title and fill it with 'NA' and maintain the exact same order:

        - Volunteering
        - Competitions
        - Conferences and Workshops
        - Tests
        - Patents
        - Scholarships
        - Extracurricular Activities
    """

    # Send 3 different queries, collect the responses, and then merge them
    query1 = part1 + resume_text
    query2 = part2 + resume_text
    query3 = part3 + resume_text

    ddgs = DDGS()

    try:
        results1 = ddgs.chat(query1, model='claude-3-haiku')
        results2 = ddgs.chat(query2, model='claude-3-haiku')
        results3 = ddgs.chat(query3, model='claude-3-haiku')
        resume_result = results1 + results2 + results3
    except Exception as e:
        st.error(f"Error: {str(e)}")
        return {}

    section_headers = {
        "Name": "Name",
        "Contact Information": "Contact Information",
        "Professional Summary": "Professional Summary",
        "Education": "Education",
        "Experience": "Experience",
        "Projects": "Projects",
        "Skills": "Skills",
        "Languages": "Languages",
        "Links": "Links",
        "Awards": "Awards",
        "Certifications": "Certifications",
        "Publications": "Publications",
        "Volunteering": "Volunteering",
        "Competitions": "Competitions",
        "Conferences and Workshops": "Conferences and Workshops",
        "Tests": "Tests",
        "Patents": "Patents",
        "Scholarships": "Scholarships",
        "Extracurricular Activities": "Extracurricular Activities"
    }

    # Initialize the dictionary with 'NA' for each field
    details = {key: "NA" for key in section_headers.keys()}

    # Sort section headers by their positions in the text
    sorted_sections = sorted([(header, resume_result.find(header)) for header in section_headers.values() if resume_result.find(header) != -1], key=lambda x: x[1])

    for i, (current_header, start_pos) in enumerate(sorted_sections):
        # Determine the end position as the start of the next header
        end_pos = sorted_sections[i + 1][1] if i + 1 < len(sorted_sections) else len(resume_result)

        # Extract content between headers
        content = resume_result[start_pos + len(current_header):end_pos].strip()

        # Map extracted content to the correct dictionary key
        for key, header in section_headers.items():
            if header == current_header:
                if "NA" in content and len(content) < 10:
                    details[key] = None
                else:
                    content = content.translate(str.maketrans('', '', '-'))
                    # Process each line to remove starting non-alphabetic characters and filter out empty lines
                    
                    cleaned_lines = []
                    for line in content.split('\n'):
                        # Remove starting non-alphabetic characters
                        line = re.sub(r'^[^a-zA-Z0-9]+', '', line)
                        if line.strip():  # Check if line is not empty after removal
                            cleaned_lines.append(line)

                    # Join the cleaned lines back together
                    cleaned_text = '\n'.join(cleaned_lines) if content else None
                    
                    details[key] = cleaned_text if content else None
                break
    return details


def generate_resume(
    name=None,
    contact_info=None,
    professional_summary=None,
    education=None,
    experience=None,
    projects=None,
    skills=None,
    languages=None,
    links=None,
    awards=None,
    certifications=None,
    publications=None,
    volunteering=None,
    competitions=None,
    conferences_workshops=None,
    tests=None,
    patents=None,
    scholarships=None,
    extracurricular_activities=None,
    output_file='resume.docx',
    output_pdf_file='resume.pdf'
):
    doc = docx.Document()

    # Add content to the DOCX file
    if name:
        doc.add_heading(name, level=0)
    if contact_info:
        doc.add_paragraph(contact_info)
    if professional_summary:
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(professional_summary)
    if education:
        doc.add_heading('Education', level=1)
        for edu in education:
            doc.add_paragraph(f"{edu['institution']} - {edu['degree']} ({edu['cgpa']}) - {edu['dates']}")
    if experience:
        doc.add_heading('Experience', level=1)
        for exp in experience:
            doc.add_paragraph(f"{exp['title']} at {exp['company']} ({exp['dates']})")
            doc.add_paragraph(exp['description'], style='List Bullet')
    if projects:
        doc.add_heading('Projects', level=1)
        for proj in projects:
            doc.add_paragraph(proj['title'])
            doc.add_paragraph(proj['description'], style='List Bullet')
    if skills:
        doc.add_heading('Skills', level=1)
        for skill in skills:
            doc.add_paragraph(f"{skill['category']}: {', '.join(skill['skills'])}")
    if languages:
        doc.add_heading('Languages', level=1)
        doc.add_paragraph(', '.join(languages))
    if links:
        doc.add_heading('Links', level=1)
        for link in links:
            doc.add_paragraph(link)
    if awards:
        doc.add_heading('Awards', level=1)
        for award in awards:
            doc.add_paragraph(award)
    if certifications:
        doc.add_heading('Certifications', level=1)
        for certification in certifications:
            doc.add_paragraph(certification)
    if publications:
        doc.add_heading('Publications', level=1)
        for publication in publications:
            doc.add_paragraph(publication)
    if volunteering:
        doc.add_heading('Volunteering', level=1)
        for vol in volunteering:
            doc.add_paragraph(vol)
    if competitions:
        doc.add_heading('Competitions', level=1)
        for competition in competitions:
            doc.add_paragraph(competition)
    if conferences_workshops:
        doc.add_heading('Conferences and Workshops', level=1)
        for conference in conferences_workshops:
            doc.add_paragraph(conference)
    if tests:
        doc.add_heading('Tests', level=1)
        for test in tests:
            doc.add_paragraph(test)
    if patents:
        doc.add_heading('Patents', level=1)
        for patent in patents:
            doc.add_paragraph(patent)
    if scholarships:
        doc.add_heading('Scholarships', level=1)
        for scholarship in scholarships:
            doc.add_paragraph(scholarship)
    if extracurricular_activities:
        doc.add_heading('Extra Curricular Activities', level=1)
        for activity in extracurricular_activities:
            doc.add_paragraph(activity)
    
    # Save the DOCX file
    doc.save(output_file)

    # Convert DOCX to PDF
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 40
    usable_width = width - 2 * margin
    usable_height = height - 2 * margin
    y_position = height - margin

    def add_text(c, text, x, y, max_width):
        lines = []
        words = text.split()
        line = ''
        for word in words:
            if c.stringWidth(line + word, 'Helvetica', 12) < max_width:
                line += (word + ' ')
            else:
                lines.append(line)
                line = word + ' '
        lines.append(line)
        
        for line in lines:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - margin
            c.drawString(x, y, line)
            y -= 14  # Line height
        return y

    def add_section_title(title, y):
        c.setFont("Helvetica-Bold", 14)
        y = add_text(c, title, margin, y, usable_width)
        y -= 10
        return y

    def add_section_content(content, y):
        c.setFont("Helvetica", 12)
        y = add_text(c, content, margin, y, usable_width)
        y -= 20
        return y

    if name:
        c.setFont("Helvetica-Bold", 16)
        y_position = add_text(c, name, margin, y_position, usable_width)
        y_position -= 30

    if contact_info:
        c.setFont("Helvetica", 12)
        y_position = add_text(c, contact_info, margin, y_position, usable_width)
        y_position -= 30

    if professional_summary:
        y_position = add_section_title("Professional Summary", y_position)
        y_position = add_section_content(professional_summary, y_position)

    if education:
        y_position = add_section_title("Education", y_position)
        for edu in education:
            y_position = add_section_content(f"{edu['institution']} - {edu['degree']} ({edu['cgpa']}) - {edu['dates']}", y_position)

    if experience:
        y_position = add_section_title("Experience", y_position)
        for exp in experience:
            y_position = add_section_content(f"{exp['title']} at {exp['company']} ({exp['dates']})", y_position)
            y_position = add_section_content(exp['description'], y_position)

    if projects:
        y_position = add_section_title("Projects", y_position)
        for proj in projects:
            y_position = add_section_content(proj['title'], y_position)
            y_position = add_section_content(proj['description'], y_position)

    if skills:
        y_position = add_section_title("Skills", y_position)
        for skill in skills:
            y_position = add_section_content(f"{skill['category']}: {', '.join(skill['skills'])}", y_position)

    if languages:
        y_position = add_section_title("Languages", y_position)
        y_position = add_section_content(', '.join(languages), y_position)

    if links:
        y_position = add_section_title("Links", y_position)
        for link in links:
            y_position = add_section_content(link, y_position)

    if awards:
        y_position = add_section_title("Awards", y_position)
        for award in awards:
            y_position = add_section_content(award, y_position)

    if certifications:
        y_position = add_section_title("Certifications", y_position)
        for certification in certifications:
            y_position = add_section_content(certification, y_position)

    if publications:
        y_position = add_section_title("Publications", y_position)
        for publication in publications:
            y_position = add_section_content(publication, y_position)

    if volunteering:
        y_position = add_section_title("Volunteering", y_position)
        for vol in volunteering:
            y_position = add_section_content(vol, y_position)

    if competitions:
        y_position = add_section_title("Competitions", y_position)
        for competition in competitions:
            y_position = add_section_content(competition, y_position)

    if conferences_workshops:
        y_position = add_section_title("Conferences and Workshops", y_position)
        for conference in conferences_workshops:
            y_position = add_section_content(conference, y_position)

    if tests:
        y_position = add_section_title("Tests", y_position)
        for test in tests:
            y_position = add_section_content(test, y_position)

    if patents:
        y_position = add_section_title("Patents", y_position)
        for patent in patents:
            y_position = add_section_content(patent, y_position)

    if scholarships:
        y_position = add_section_title("Scholarships", y_position)
        for scholarship in scholarships:
            y_position = add_section_content(scholarship, y_position)

    if extracurricular_activities:
        y_position = add_section_title("Extra Curricular Activities", y_position)
        for activity in extracurricular_activities:
            y_position = add_section_content(activity, y_position)

    # Finish the PDF
    c.save()
    
    with open(output_pdf_file, 'wb') as f:
        f.write(buffer.getvalue())

    return output_pdf_file

    
    
# Function to estimate text area height based on content length
def estimate_height(text):
    estimated_height = int(len(text) / 1.8)  # Adjust the divisor to control height increments
    return 300 if estimated_height > 300 else estimated_height

# Streamlit app code
def main():
    st.title("Resume Builder")

    # Step 1: Ask user if they have a current resume
    has_resume = st.radio("Do you have a current resume?", ("Yes", "No"))

    # Dictionary to store user inputs
    fields = {
        "Name": "",
        "Contact Information": "",
        "Professional Summary": "",
        "Education": "",
        "Experience": "",
        "Projects": "",
        "Skills": "",
        "Languages": "",
        "Links": "",
        "Awards": "",
        "Certifications": "",
        "Publications": "",
        "Volunteering": "",
        "Competitions": "",
        "Conferences and Workshops": "",
        "Tests": "",
        "Patents": "",
        "Scholarships": "",
        "Extracurricular Activities": ""
    }

    # If user has a resume, ask for upload and extract details
    if has_resume == "Yes":
        uploaded_file = st.file_uploader("Upload your current resume (PDF)", type="pdf")

        if uploaded_file is not None:
            resume_text = extract_text_from_pdf(uploaded_file)

            # Extract sections using AI
            extracted_sections = extract_sections_from_resume(resume_text)

            # Update fields dictionary with extracted details
            for key in fields:
                fields[key] = extracted_sections.get(key, "")

            st.success("Resume text extracted successfully!")
        else:
            st.warning("Please upload a PDF file.")

    # Step 2: Display fields and get user inputs
    st.header("Fill in or Edit Your Resume Information")
    for key in fields:
        content = fields[key]
        height = estimate_height(content) if content is not None else 0
        fields[key] = st.text_area(key, value=content, height=height)

    # Preview resume on the right side
    with st.expander("Preview Resume", expanded=False):
        preview_content = generate_resume_markdown(
            name = fields["Name"] if fields["Name"] else None,
            contact_info = fields["Contact Information"] if fields["Contact Information"] else None,
            professional_summary = fields["Professional Summary"] if fields["Professional Summary"] else None,
            education = (
                [{'institution': parts[0], 'degree': parts[1] if len(parts) > 1 else "", 'cgpa': parts[2] if len(parts) > 2 else "", 'dates': parts[3] if len(parts) > 3 else ""}
                for x in fields["Education"].split('\n') if x and (parts := x.split(' - ')) and len(parts) >= 1]
                if fields["Education"] else None
            ),
            experience = (
                [{'title': parts_at[0], 'company': parts_at[1].split(' (')[0] if len(parts_at) > 1 else "", 'dates': parts_paren[1][:-1] if len(parts_paren) > 1 else "", 'description': fields["Experience"].split('\n')[1] if len(fields["Experience"].split('\n')) > 1 else ""}
                for x in fields["Experience"].split('\n') if x and (parts_at := x.split(' at ')) and (parts_paren := x.split(' (')) and len(parts_at) >= 1]
                if fields["Experience"] else None
            ),
            projects = (
                [{'title': parts[0], 'description': parts[1] if len(parts) > 1 else ""}
                for x in fields["Projects"].split('\n\n') if (parts := x.split('\n')) and len(parts) > 0]
                if fields["Projects"] else None
            ),
            skills = (
                [{'category': 'General Skills', 'skills': fields["Skills"].split(', ')}]
                if fields["Skills"] else None
            ),
            languages = fields["Languages"].split('\n') if fields["Languages"] else None,
            links = fields["Links"].split('\n') if fields["Links"] else None,
            awards = fields["Awards"].split('\n') if fields["Awards"] else None,
            certifications = fields["Certifications"].split('\n') if fields["Certifications"] else None,
            publications = fields["Publications"].split('\n') if fields["Publications"] else None,
            volunteering = fields["Volunteering"].split('\n') if fields["Volunteering"] else None,
            competitions = fields["Competitions"].split('\n') if fields["Competitions"] else None,
            conferences_workshops = fields["Conferences and Workshops"].split('\n') if fields["Conferences and Workshops"] else None,
            tests = fields["Tests"].split('\n') if fields["Tests"] else None,
            patents = fields["Patents"].split('\n') if fields["Patents"] else None,
            scholarships = fields["Scholarships"].split('\n') if fields["Scholarships"] else None,
            extracurricular_activities = fields["Extracurricular Activities"].split('\n') if fields["Extracurricular Activities"] else None

        )
        st.markdown(preview_content)

    # Step 3: Generate Resume on button click
    if st.button("Generate Resume"):
        pdf_file = generate_resume(
            name=fields["Name"] if fields["Name"] else None,
            contact_info=fields["Contact Information"] if fields["Contact Information"] else None,
            professional_summary=fields["Professional Summary"] if fields["Professional Summary"] else None,
            education=[{'institution': parts[0], 'degree': parts[1] if len(parts) > 1 else "", 'cgpa': parts[2] if len(parts) > 2 else "", 'dates': parts[3] if len(parts) > 3 else ""} for x in fields["Education"].split('\n') if x and (parts := x.split(' - ')) and len(parts) >= 1] if fields["Education"] else None,
            experience=[{'title': parts_at[0], 'company': parts_at[1].split(' (')[0] if len(parts_at) > 1 else "", 'dates': parts_paren[1][:-1] if len(parts_paren) > 1 else "", 'description': fields["Experience"].split('\n')[1] if len(fields["Experience"].split('\n')) > 1 else ""} for x in fields["Experience"].split('\n') if x and (parts_at := x.split(' at ')) and (parts_paren := x.split(' (')) and len(parts_at) >= 1] if fields["Experience"] else None,
            projects=[{'title': parts[0], 'description': parts[1] if len(parts) > 1 else ""} for x in fields["Projects"].split('\n\n') if (parts := x.split('\n')) and len(parts) > 0] if fields["Projects"] else None,
            skills=[{'category': 'General Skills', 'skills': fields["Skills"].split(', ')}] if fields["Skills"] else None,
            languages=fields["Languages"].split('\n') if fields["Languages"] else None,
            links=fields["Links"].split('\n') if fields["Links"] else None,
            awards=fields["Awards"].split('\n') if fields["Awards"] else None,
            certifications=fields["Certifications"].split('\n') if fields["Certifications"] else None,
            publications=fields["Publications"].split('\n') if fields["Publications"] else None,
            volunteering=fields["Volunteering"].split('\n') if fields["Volunteering"] else None,
            competitions=fields["Competitions"].split('\n') if fields["Competitions"] else None,
            conferences_workshops=fields["Conferences and Workshops"].split('\n') if fields["Conferences and Workshops"] else None,
            tests=fields["Tests"].split('\n') if fields["Tests"] else None,
            patents=fields["Patents"].split('\n') if fields["Patents"] else None,
            scholarships=fields["Scholarships"].split('\n') if fields["Scholarships"] else None,
            extracurricular_activities=fields["Extracurricular Activities"].split('\n') if fields["Extracurricular Activities"] else None
        )
        
        # Download link for the PDF
        with open(pdf_file, "rb") as f:
            st.download_button("Download PDF", f, file_name="resume.pdf")

if __name__ == "__main__":
    main()
